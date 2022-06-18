from docx import Document
import re
from docx2pdf import convert
document = Document('schedule.docx')
table = document.tables[0]

data = []

keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)
    if i == 0:
        keys = tuple(text)
        continue
    row_data = dict(zip(keys, text))
    data.append(row_data)

text = ''
for row in table.rows:
    for cell in row.cells:
       for para in cell.paragraphs:
           text+=para.text+" "
pattern = re.compile(r'[\d]+\:[\d]+[\s]*AM')
matches = pattern.finditer(text)
ams = []
for match in matches:
    ams.append(match.group())
for i in range(len(ams)):
    ams[i] = ams[i].replace(' AM', '')
    ams[i] = ams[i].replace('AM', '')
pattern = re.compile(r'[\d]+\:[\d]+[\s]*PM')
matches = pattern.finditer(text)
pms = []
for match in matches:
    pms.append(match.group())
for i in range(len(pms)):
    pms[i] = pms[i].replace(' PM', '')
    pms[i] = pms[i].replace('PM', '')
    if pms[i].startswith('1:'):
        pms[i] = pms[i].replace('1:', '13:')
    if pms[i].startswith('2:'):
        pms[i] = pms[i].replace('2:', '14:')
    if pms[i].startswith('3:'):
        pms[i] = pms[i].replace('3:', '15:')
    if pms[i].startswith('4:'):
        pms[i] = pms[i].replace('4:', '16:')
    if pms[i].startswith('5:'):
        pms[i] = pms[i].replace('5:', '17:')
    if pms[i].startswith('6:'):
        pms[i] = pms[i].replace('6:', '18:')
    if pms[i].startswith('7:'):
        pms[i] = pms[i].replace('7:', '19:')
    if pms[i].startswith('8:'):
        pms[i] = pms[i].replace('8:', '20:')
    if pms[i].startswith('9:'):
        pms[i] = pms[i].replace('9:', '21:')
    if pms[i].startswith('10:'):
        pms[i] = pms[i].replace('10:', '22:')
    if pms[i].startswith('11:'):
        pms[i] = pms[i].replace('11:', '23:')
times = ams + pms
j = 0
for i in range(len(data)):
    data[i]["Time"] = times[j]
    if j+1 < len(times):
        data[i]["Time"] += " - "+times[j + 1]
    j = j + 2

doc = Document()
table = doc.add_table(rows=1, cols=2)
row = table.rows[0].cells
row[0].text = 'Time'
row[1].text = 'Activity'
for d in data:
    row = table.add_row().cells
    row[0].text = d["Time"]
    row[1].text = d["Activity"]
table.style = "Table Grid"
doc.save('new-schedule.docx')
convert('new-schedule.docx')